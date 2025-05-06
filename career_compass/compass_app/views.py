from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from django.http import HttpResponse, FileResponse
from django.core.mail import send_mail
from django.conf import settings
from .models import (
    UserProfile, Resume, Skill, CareerPath, CareerAnalysis,
    SkillMatch, RecommendedSkill, CareerPathRecommendation,
    JobPortal, JobRecommendation
)
import random  # Just for demo purposes
import os
import io
import platform
import subprocess
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tempfile
from django.conf import settings
from docx2pdf import convert
import datetime
import PyPDF2
import re
from django.views.decorators.http import require_POST

# Cached result of Word installation check
_word_installed_cache = None

# Helper function to check if Microsoft Word is installed
def is_word_installed(force_check=False):
    """Check if Microsoft Word is installed on the system
    
    Args:
        force_check: If True, ignore the cached result and check again
        
    Returns:
        bool: True if Word is installed, False otherwise
    """
    global _word_installed_cache
    
    # Return cached result if available
    if _word_installed_cache is not None and not force_check:
        return _word_installed_cache
    
    system = platform.system()
    
    if system == 'Windows':
        # Check common Word installation paths on Windows
        common_paths = [
            r"C:\Program Files\Microsoft Office\root\Office16\WINWORD.EXE",
            r"C:\Program Files (x86)\Microsoft Office\root\Office16\WINWORD.EXE",
            r"C:\Program Files\Microsoft Office\Office16\WINWORD.EXE",
            r"C:\Program Files (x86)\Microsoft Office\Office16\WINWORD.EXE",
            r"C:\Program Files\Microsoft Office\Office15\WINWORD.EXE",
            r"C:\Program Files (x86)\Microsoft Office\Office15\WINWORD.EXE",
            r"C:\Program Files\Microsoft Office\Office14\WINWORD.EXE",
            r"C:\Program Files (x86)\Microsoft Office\Office14\WINWORD.EXE",
        ]
        
        for path in common_paths:
            if os.path.exists(path):
                _word_installed_cache = True
                return True
                
        # Also check using Windows registry
        try:
            import winreg
            key = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\WINWORD.EXE")
            winreg.CloseKey(key)
            _word_installed_cache = True
            return True
        except:
            pass
            
    elif system == 'Darwin':  # macOS
        try:
            result = subprocess.run(['mdfind', 'kMDItemCFBundleIdentifier == "com.microsoft.Word"'], 
                                  capture_output=True, text=True, check=True)
            has_word = bool(result.stdout.strip())
            _word_installed_cache = has_word
            return has_word
        except:
            pass
    
    # For Linux or as a fallback, check if the 'convert' command from docx2pdf works
    try:
        # Create a small test document
        test_doc = Document()
        test_doc.add_paragraph("Test")
        test_docx = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        test_doc.save(test_docx.name)
        test_docx.close()
        
        # Try to convert it (will fail quickly if no Word)
        test_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        test_pdf.close()
        
        # Use a short timeout to avoid hanging
        subprocess.run(['python', '-c', 
                      f"from docx2pdf import convert; convert('{test_docx.name}', '{test_pdf.name}')"],
                     timeout=3)
        
        # Clean up
        os.unlink(test_docx.name)
        if os.path.exists(test_pdf.name):
            os.unlink(test_pdf.name)
            _word_installed_cache = True
            return True
        
        _word_installed_cache = False
        return False
    except Exception as e:
        print(f"Word check failed: {str(e)}")
        _word_installed_cache = False
        return False
    
    _word_installed_cache = False
    return False

# Helper function to check if user is admin
def is_admin(user):
    return user.is_superuser

# User views
def home(request):
    job_portals = JobPortal.objects.all()
    return render(request, 'compass_app/home.html', {'job_portals': job_portals})

def dashboard(request):
    # Create a default user for demo purposes if no authentication
    user = get_default_user()
    
    # Get analysis data from session if available
    resume_processed = request.session.get('resume_processed', False)
    analysis_data = request.session.get('resume_analysis', {})
    
    # Calculate counts for dashboard stats
    context = {
        'resume_processed': resume_processed,
        'resume_count': 1 if resume_processed else 0,
        'job_count': len(analysis_data.get('job_recommendations', [])) if resume_processed else 0,
        'skill_count': len(analysis_data.get('skill_matches', [])) if resume_processed else 0,
        'career_count': len(analysis_data.get('career_recommendations', [])) if resume_processed else 0,
        'ats_score': analysis_data.get('ats_score', 0) if resume_processed else None,
        'resume_filename': request.session.get('resume_filename', '') if resume_processed else '',
        'resume_uploaded_at': request.session.get('resume_uploaded_at', '') if resume_processed else ''
    }
    
    return render(request, 'compass_app/dashboard.html', context)

def resume_upload(request):
    # Create a default user for demo purposes if no authentication
    user = get_default_user()
    
    if request.method == 'POST' and request.FILES.get('resume'):
        resume_file = request.FILES['resume']
        
        # Process the resume directly without saving it
        file_extension = resume_file.name.split('.')[-1].lower()
        
        # Generate analysis data and store in session
        analysis_data = generate_resume_analysis(resume_file, user)
        
        # Store analysis data in session
        request.session['resume_analysis'] = analysis_data
        request.session['resume_processed'] = True
        request.session['resume_filename'] = resume_file.name
        request.session['resume_uploaded_at'] = str(datetime.datetime.now())
        
        messages.success(request, 'Resume uploaded successfully! Analyzing...')
        return redirect('career_analysis')
    
    return render(request, 'compass_app/resume_upload.html')

def generate_resume_analysis(resume_file, user):
    """Generate analysis data from a resume file and attempt to extract actual user information"""
    
    # Extract text content from the resume file based on its type
    resume_content = ""
    file_extension = resume_file.name.split('.')[-1].lower()
    
    try:
        if file_extension == 'pdf':
            import io
            import PyPDF2
            
            # Read the PDF file content
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(resume_file.read()))
            for page in range(len(pdf_reader.pages)):
                page_text = pdf_reader.pages[page].extract_text() or ""
                resume_content += page_text + "\n"
            
            # Reset file cursor for future operations
            resume_file.seek(0)
            
        elif file_extension in ['docx', 'doc']:
            from docx import Document
            import io
            
            # Read the DOCX file content
            doc = Document(io.BytesIO(resume_file.read()))
            for para in doc.paragraphs:
                if para.text:
                    resume_content += para.text + "\n"
                    
            # Also check tables for information (common in resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            resume_content += cell.text + "\n"
            
            # Reset file cursor for future operations
            resume_file.seek(0)
            
        else:
            # For other file types, try to read as plain text
            content = resume_file.read()
            try:
                resume_content = content.decode('utf-8', errors='ignore')
            except (UnicodeDecodeError, AttributeError):
                resume_content = str(content)
            resume_file.seek(0)
    
    except Exception as e:
        print(f"Error extracting resume content: {str(e)}")
        resume_content = ""
    
    # Debug: Print the first 500 characters of the resume content
    print(f"Resume content preview: {resume_content[:500]}")
    
    # Extract user information from resume content
    import re
    
    # Try to extract email with a more robust pattern
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    email_matches = re.findall(email_pattern, resume_content)
    email = email_matches[0] if email_matches else user.email
    
    # Try to extract phone number (various formats)
    phone_patterns = [
        r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # 123-456-7890, 123.456.7890, 123 456 7890
        r'\b\(\d{3}\)[-.\s]?\d{3}[-.\s]?\d{4}\b',  # (123)-456-7890, (123).456.7890, (123) 456 7890
        r'\b\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',  # +1-123-456-7890, +91 123 456 7890
        r'\b\d{10}\b'  # 1234567890
    ]
    
    phone = getattr(user.profile, 'phone', '1234567890')
    for pattern in phone_patterns:
        phone_matches = re.findall(pattern, resume_content)
        if phone_matches:
            phone = phone_matches[0]
            break
    
    # Try to extract name using multiple approaches
    # Method 1: Look at the first few lines (common resume format)
    name = user.get_full_name() or user.username
    first_lines = resume_content.strip().split('\n')[:5]
    for line in first_lines:
        # Look for a line with just a name (typically the first line of a resume)
        line = line.strip()
        if line and 3 <= len(line.split()) <= 4 and all(word[0].isupper() for word in line.split()):
            name = line
            break
    
    # Method 2: If still using default, try regex pattern for common name formats
    if name == user.get_full_name() or name == user.username:
        name_patterns = [
            r'^([A-Z][a-z]+(?: [A-Z][a-z]+)+)',  # FirstName LastName at start of line
            r'\b([A-Z][a-z]+ [A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b'  # FirstName LastName or FirstName MiddleName LastName
        ]
        
        for pattern in name_patterns:
            name_matches = re.findall(pattern, resume_content)
            if name_matches:
                name = name_matches[0]
                break
    
    # Try to extract education information with improved patterns
    education = getattr(user.profile, 'education', 'Bachelor of Technology')
    education_patterns = [
        r'(?:Bachelors?|B\.?(?:Tech|Sc|A|B|E))\s+(?:of|in)?\s+([A-Za-z\s&]+)',
        r'(?:Masters?|M\.?(?:Tech|Sc|A|B|E))\s+(?:of|in)?\s+([A-Za-z\s&]+)',
        r'(?:Doctorate|Ph\.?D\.?)\s+(?:of|in)?\s+([A-Za-z\s&]+)',
        r'(?:Diploma)\s+(?:in)?\s+([A-Za-z\s&]+)'
    ]
    
    for pattern in education_patterns:
        matches = re.findall(pattern, resume_content, re.IGNORECASE)
        if matches:
            found_education = matches[0]
            if len(found_education.strip()) > 3:  # Ensure it's a meaningful match
                education = f"{re.search(pattern, resume_content, re.IGNORECASE).group(0)}"
                break
    
    # Try to extract university/institution
    university = getattr(user.profile, 'university', 'Demo University')
    university_patterns = [
        r'(?:University|College|Institute|School)\s+of\s+([A-Za-z\s]+)',
        r'([A-Za-z]+\s+(?:University|College|Institute|School))',
        r'([A-Za-z]+\s+Institute\s+of\s+[A-Za-z\s]+)',
        r'([A-Za-z]+\s+College\s+of\s+[A-Za-z\s]+)'
    ]
    
    for pattern in university_patterns:
        matches = re.findall(pattern, resume_content, re.IGNORECASE)
        if matches:
            for match in matches:
                if len(match.strip()) > 3:  # Ensure it's a meaningful match
                    university = re.search(pattern, resume_content, re.IGNORECASE).group(0).title()
                    break
    
    # Try to extract graduation year (looking for 4-digit years, typically recent)
    graduation_year = getattr(user.profile, 'graduation_year', 2023)
    current_year = datetime.datetime.now().year
    
    # Look for years in context of education
    edu_year_patterns = [
        r'(?:graduate|graduation|completed|class of|batch|passing|passed out|year).*?(\d{4})',
        r'(?:education|university|college|school|degree).*?(\d{4})',
        r'(\d{4}).*?(?:graduate|graduation|completed|present)'
    ]
    
    # Find all 4-digit years in the resume
    all_years = re.findall(r'\b((?:19|20)\d{2})\b', resume_content)
    possible_grad_years = [int(y) for y in all_years if 1990 <= int(y) <= current_year + 4]
    
    # First try education-specific contexts
    for pattern in edu_year_patterns:
        matches = re.findall(pattern, resume_content.lower())
        if matches:
            try:
                year = int(matches[0])
                if 1990 <= year <= current_year + 4:  # Reasonable range for graduation year
                    graduation_year = year
                    break
            except:
                pass
    
    # If not found, use the most recent year from all years found
    if graduation_year == 2023 and possible_grad_years:
        graduation_year = max(possible_grad_years)
    
    # Parse the content to extract skills
    common_skills = [
        "Python", "Java", "JavaScript", "HTML", "CSS", "React", "Angular", "Node.js", 
        "Django", "Flask", "Express", "SQL", "MySQL", "PostgreSQL", "MongoDB", 
        "Git", "Docker", "AWS", "Azure", "GCP", "Linux", "Windows", "MacOS", 
        "C++", "C#", "Ruby", "PHP", "Go", "Swift", "Kotlin", "TypeScript",
        "Communication", "Teamwork", "Leadership", "Problem Solving", "Critical Thinking",
        "Time Management", "Project Management", "Agile", "Scrum", "DevOps",
        ".NET", "Spring", "Hibernate", "REST API", "GraphQL", "JSON", "XML", "YAML",
        "NoSQL", "Jenkins", "Kubernetes", "CI/CD", "Machine Learning", "Data Science",
        "TensorFlow", "PyTorch", "NLP", "Computer Vision", "Artificial Intelligence"
    ]
    
    # Create a list of skills that appear in the resume
    skills = []
    for skill in common_skills:
        if re.search(r'\b' + re.escape(skill) + r'\b', resume_content, re.IGNORECASE):
            skills.append(skill)
    
    # Enhanced ATS score calculation based on multiple factors
    ats_score_factors = {
        'has_contact_info': 0,  # Max 15 points
        'has_education': 0,     # Max 15 points
        'has_skills': 0,        # Max 25 points
        'has_experience': 0,    # Max 20 points
        'formatting': 0,        # Max 15 points
        'keyword_density': 0    # Max 10 points
    }
    
    # Check contact info completeness (email, phone)
    if re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', resume_content):
        ats_score_factors['has_contact_info'] += 8
    
    # Check for phone number
    if re.search(r'\b(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b', resume_content):
        ats_score_factors['has_contact_info'] += 7
    
    # Check education details
    edu_terms = ['degree', 'bachelor', 'master', 'phd', 'diploma', 'certification', 'university', 'college', 'institute', 'school']
    education_count = sum(1 for term in edu_terms if re.search(r'\b' + re.escape(term) + r'\b', resume_content, re.IGNORECASE))
    ats_score_factors['has_education'] = min(15, education_count * 3)
    
    # Check skills (based on matched skills)
    skill_count = len(skills)
    ats_score_factors['has_skills'] = min(25, skill_count * 3)
    
    # Check experience indicators
    exp_terms = ['experience', 'work', 'job', 'position', 'role', 'employment', 'career', 'years']
    experience_count = sum(1 for term in exp_terms if re.search(r'\b' + re.escape(term) + r'\b', resume_content, re.IGNORECASE))
    ats_score_factors['has_experience'] = min(20, experience_count * 2.5)
    
    # Check formatting (based on structure)
    format_score = 10  # Base score
    
    # Check for sections with headers (typically indicate good formatting)
    section_headers = ['experience', 'education', 'skills', 'summary', 'objective', 'profile', 'projects', 'certifications']
    section_count = sum(1 for header in section_headers if re.search(r'\b' + re.escape(header) + r'\b', resume_content, re.IGNORECASE))
    format_score += min(5, section_count)
    
    ats_score_factors['formatting'] = format_score
    
    # Calculate keyword density for relevant terms
    # Count occurrences of industry terminology
    industry_terms = ['develop', 'manage', 'lead', 'create', 'design', 'implement', 'analyze', 'coordinate', 'improve', 'increase']
    keyword_count = sum(len(re.findall(r'\b' + re.escape(term) + r'\b', resume_content, re.IGNORECASE)) for term in industry_terms)
    
    # Content length matters for keyword density
    content_length = len(resume_content.split())
    if content_length > 0:
        keyword_density = (keyword_count / content_length) * 100
        ats_score_factors['keyword_density'] = min(10, keyword_density * 2)
    
    # Calculate total ATS score (0-100)
    ats_score = sum(ats_score_factors.values())
    
    # Ensure the score is between 0 and 100
    ats_score = max(0, min(100, int(ats_score)))
    
    # Extract user profile information 
    user_profile = {
        'username': user.username,
        'email': email,
        'first_name': name.split()[0] if ' ' in name else name,
        'last_name': name.split()[-1] if ' ' in name else '',
        'full_name': name,
        'profile': {
            'education': education,
            'university': university,
            'graduation_year': graduation_year,
            'phone': phone
        }
    }
    
    # Debug: Print the extracted profile information
    print(f"Extracted user profile: {user_profile}")
    
    # Create skills with random proficiency values for skills found in the resume
    skill_matches = []
    for skill in skills[:7]:  # Limit to 7 skills for display
        skill_matches.append({
            'name': skill,
            'proficiency': random.randint(50, 95)
        })
    
    # If no skills were found, use some default skills
    if not skill_matches:
        skill_matches = [
            {'name': 'Python', 'proficiency': random.randint(50, 95)},
            {'name': 'Java', 'proficiency': random.randint(50, 95)},
            {'name': 'JavaScript', 'proficiency': random.randint(50, 95)},
            {'name': 'HTML', 'proficiency': random.randint(50, 95)},
            {'name': 'CSS', 'proficiency': random.randint(50, 95)},
            {'name': 'Communication', 'proficiency': random.randint(50, 95)},
            {'name': 'Teamwork', 'proficiency': random.randint(50, 95)}
        ]
    
    # Generate some strengths and areas to improve based on extracted skills
    strengths = []
    areas_to_improve = []
    
    # Technical skills as strengths
    tech_skills = ["Python", "Java", "JavaScript", "HTML", "CSS", "React", "Angular", 
                  "SQL", "Git", "Docker", "AWS", "Azure"]
    # Soft skills as strengths
    soft_skills = ["Communication", "Teamwork", "Leadership", "Problem Solving"]
    
    # Advanced skills as areas to improve
    advanced_skills = ["Machine Learning", "Data Science", "TensorFlow", "DevOps", 
                      "Kubernetes", "GraphQL", "Microservices", "System Design"]
    
    # Add skills found in resume to strengths
    for skill in skills:
        if skill in tech_skills or skill in soft_skills:
            strengths.append(skill)
        elif skill in advanced_skills:
            areas_to_improve.append(skill)
    
    # Ensure we have at least 3 strengths and areas to improve
    default_strengths = ["Communication", "Problem-solving", "Teamwork"]
    default_areas = ["Technical knowledge", "Leadership", "Time management"]
    
    strengths = strengths[:3] or default_strengths
    areas_to_improve = areas_to_improve[:3] or default_areas
    
    # Mock analysis data with extracted user info and skills
    analysis_data = {
        'user_id': user.id,
        'user_info': user_profile,  # Add extracted user profile info to session data
        'strengths': ", ".join(strengths),
        'areas_to_improve': ", ".join(areas_to_improve),
        'ats_score': ats_score,
        'created_at': str(datetime.datetime.now()),
        
        # Use extracted skills with random proficiency levels
        'skill_matches': skill_matches,
        
        # Generate recommended skills (that aren't already in skill_matches)
        'recommended_skills': [
            {'name': 'Django', 'importance': random.randint(70, 95)},
            {'name': 'React', 'importance': random.randint(70, 95)},
            {'name': 'Node.js', 'importance': random.randint(70, 95)},
            {'name': 'Data Analysis', 'importance': random.randint(70, 95)},
            {'name': 'Machine Learning', 'importance': random.randint(70, 95)}
        ],
        
        # Generate career path recommendations
        'career_recommendations': [
            {
                'title': 'Software Developer',
                'description': 'Develop applications and systems using programming languages.',
                'avg_salary': '₹6-15 LPA',
                'growth_potential': 'High',
                'match_percentage': random.randint(60, 95)
            },
            {
                'title': 'Data Analyst',
                'description': 'Analyze and interpret complex data to inform business decisions.',
                'avg_salary': '₹5-12 LPA',
                'growth_potential': 'High',
                'match_percentage': random.randint(60, 95)
            },
            {
                'title': 'Web Developer', 
                'description': 'Create and maintain websites and web applications.',
                'avg_salary': '₹4-10 LPA',
                'growth_potential': 'Medium',
                'match_percentage': random.randint(60, 95)
            },
            {
                'title': 'UI/UX Designer',
                'description': 'Design user interfaces and experiences for digital products.',
                'avg_salary': '₹4-8 LPA',
                'growth_potential': 'Medium',
                'match_percentage': random.randint(60, 95)
            },
            {
                'title': 'Project Manager',
                'description': 'Plan, execute, and close projects within constraints.',
                'avg_salary': '₹8-18 LPA',
                'growth_potential': 'Medium',
                'match_percentage': random.randint(60, 95)
            }
        ],
        
        # Generate job recommendations
        'job_recommendations': [
            {
                'title': 'Junior Software Developer',
                'company': 'TCS',
                'description': 'Entry-level development role',
                'portal': 'Naukri.com',
                'url': 'https://www.naukri.com/job-listings-junior-software-developer'
            },
            {
                'title': 'Graduate Trainee',
                'company': 'Infosys',
                'description': 'Training program for fresh graduates',
                'portal': 'Indeed.com',
                'url': 'https://www.indeed.co.in/viewjob?jk=12345'
            },
            {
                'title': 'Web Developer Intern',
                'company': 'Wipro',
                'description': 'Web development internship opportunity',
                'portal': 'LinkedIn',
                'url': 'https://www.linkedin.com/jobs/view/12345'
            },
            {
                'title': 'Associate Data Analyst',
                'company': 'Cognizant',
                'description': 'Entry-level data analysis position',
                'portal': 'Naukri.com',
                'url': 'https://www.naukri.com/job-listings-associate-data-analyst'
            },
            {
                'title': 'UI Designer',
                'company': 'Tech Mahindra',
                'description': 'Junior UI designer role',
                'portal': 'Indeed.com',
                'url': 'https://www.indeed.co.in/viewjob?jk=67890'
            }
        ]
    }
    
    return analysis_data

def career_analysis(request):
    # Check if resume has been processed from session
    if not request.session.get('resume_processed', False):
        messages.info(request, 'Please upload your resume first.')
        return redirect('resume_upload')
    
    # Get analysis data from session
    analysis_data = request.session.get('resume_analysis', {})
    
    # Get user information from analysis data or use default user as fallback
    user_info = analysis_data.get('user_info', None)
    if not user_info:
        # Fallback to default user if user_info is not available
        user = get_default_user()
        user_info = {
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'full_name': user.get_full_name() or user.username,
            'profile': {
                'education': getattr(user.profile, 'education', ''),
                'university': getattr(user.profile, 'university', ''),
                'graduation_year': getattr(user.profile, 'graduation_year', ''),
                'phone': getattr(user.profile, 'phone', '')
            }
        }
    
    context = {
        'user_info': user_info,
        'strengths': analysis_data.get('strengths', ''),
        'areas_to_improve': analysis_data.get('areas_to_improve', ''),
        'ats_score': analysis_data.get('ats_score', 0),
        'career_recommendations': analysis_data.get('career_recommendations', []),
        'resume_filename': request.session.get('resume_filename', 'Resume'),
        'resume_uploaded_at': request.session.get('resume_uploaded_at', '')
    }
    
    return render(request, 'compass_app/career_analysis.html', context)

def job_recommendations(request):
    # Check if resume has been processed from session
    if not request.session.get('resume_processed', False):
        messages.info(request, 'Please upload your resume first.')
        return redirect('resume_upload')
    
    # Get analysis data from session
    analysis_data = request.session.get('resume_analysis', {})
    
    # Get user information from analysis data or use default user as fallback
    user_info = analysis_data.get('user_info', None)
    if not user_info:
        # Fallback to default user if user_info is not available
        user = get_default_user()
        user_info = {
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'full_name': user.get_full_name() or user.username,
            'profile': {
                'education': getattr(user.profile, 'education', ''),
                'university': getattr(user.profile, 'university', ''),
                'graduation_year': getattr(user.profile, 'graduation_year', ''),
                'phone': getattr(user.profile, 'phone', '')
            }
        }
    
    context = {
        'user_info': user_info,
        'ats_score': analysis_data.get('ats_score', 0),
        'job_recommendations': analysis_data.get('job_recommendations', []),
        'resume_filename': request.session.get('resume_filename', 'Resume'),
        'resume_uploaded_at': request.session.get('resume_uploaded_at', '')
    }
    
    return render(request, 'compass_app/job_recommendations.html', context)

def skill_analysis(request):
    # Check if resume has been processed from session
    if not request.session.get('resume_processed', False):
        messages.info(request, 'Please upload your resume first.')
        return redirect('resume_upload')
    
    # Get analysis data from session
    analysis_data = request.session.get('resume_analysis', {})
    
    # Get user information from analysis data or use default user as fallback
    user_info = analysis_data.get('user_info', None)
    if not user_info:
        # Fallback to default user if user_info is not available
        user = get_default_user()
        user_info = {
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'full_name': user.get_full_name() or user.username,
            'profile': {
                'education': getattr(user.profile, 'education', ''),
                'university': getattr(user.profile, 'university', ''),
                'graduation_year': getattr(user.profile, 'graduation_year', ''),
                'phone': getattr(user.profile, 'phone', '')
            }
        }
    
    context = {
        'user_info': user_info,
        'skill_matches': analysis_data.get('skill_matches', []),
        'recommended_skills': analysis_data.get('recommended_skills', []),
        'ats_score': analysis_data.get('ats_score', 0),
        'resume_filename': request.session.get('resume_filename', 'Resume'),
        'resume_uploaded_at': request.session.get('resume_uploaded_at', '')
    }
    
    return render(request, 'compass_app/skill_analysis.html', context)

def print_report(request, analysis_id=None):
    # Check if resume has been processed from session
    if not request.session.get('resume_processed', False):
        messages.info(request, 'Please upload your resume first.')
        return redirect('resume_upload')
    
    # Get analysis data from session
    analysis_data = request.session.get('resume_analysis', {})
    
    # Get user information from analysis data or use default user as fallback
    user_info = analysis_data.get('user_info', None)
    if not user_info:
        # Fallback to default user if user_info is not available
        user = get_default_user()
        user_info = {
            'username': user.username,
            'email': user.email,
            'first_name': user.first_name,
            'last_name': user.last_name,
            'full_name': user.get_full_name() or user.username,
            'profile': {
                'education': getattr(user.profile, 'education', ''),
                'university': getattr(user.profile, 'university', ''),
                'graduation_year': getattr(user.profile, 'graduation_year', ''),
                'phone': getattr(user.profile, 'phone', '')
            }
        }
    
    # Create a context dictionary with data directly from session
    context = {
        'user_info': user_info,
        'report_date': datetime.datetime.now().strftime('%d %b %Y, %I:%M %p'),
        'strengths': analysis_data.get('strengths', ''),
        'areas_to_improve': analysis_data.get('areas_to_improve', ''),
        'ats_score': analysis_data.get('ats_score', 0),
        'skill_matches': analysis_data.get('skill_matches', []),
        'recommended_skills': analysis_data.get('recommended_skills', []),
        'career_recommendations': analysis_data.get('career_recommendations', []),
        'job_recommendations': analysis_data.get('job_recommendations', []),
        'resume_filename': request.session.get('resume_filename', 'Resume'),
        'resume_uploaded_at': request.session.get('resume_uploaded_at', '')
    }
    
    return render(request, 'compass_app/print_report.html', context)

# Admin views
@user_passes_test(is_admin)
def admin_dashboard(request):
    total_users = User.objects.filter(is_superuser=False).count()
    total_resumes = Resume.objects.count()
    total_analyses = CareerAnalysis.objects.count()
    
    context = {
        'total_users': total_users,
        'total_resumes': total_resumes,
        'total_analyses': total_analyses,
    }
    return render(request, 'compass_app/admin/dashboard.html', context)

@user_passes_test(is_admin)
def admin_users(request):
    users = User.objects.filter(is_superuser=False)
    context = {
        'users': users
    }
    return render(request, 'compass_app/admin/users.html', context)

@user_passes_test(is_admin)
def admin_analytics(request):
    # This would be more complex in a real app
    career_paths = CareerPath.objects.all()
    career_path_counts = []
    
    for path in career_paths:
        count = CareerPathRecommendation.objects.filter(career_path=path).count()
        career_path_counts.append({
            'path': path.title,
            'count': count
        })
    
    context = {
        'career_path_counts': career_path_counts
    }
    return render(request, 'compass_app/admin/analytics.html', context)

# Helper function to get a default user for demo purposes
def get_default_user():
    """Get or create a default demo user"""
    username = "demo_user"
    user, created = User.objects.get_or_create(
        username=username,
        defaults={
            'email': 'demo@example.com',
            'first_name': 'Demo',
            'last_name': 'User',
        }
    )
    
    # Create user profile if needed
    if created:
        UserProfile.objects.create(
            user=user,
            education='Bachelor of Technology',
            university='Demo University',
            graduation_year=2023,
            phone='1234567890'
        )
    
    return user

# Helper function to create mock analysis (for demo purposes)
def create_mock_analysis(resume):
    # Create a mock analysis
    analysis = CareerAnalysis.objects.create(
        user=resume.user,
        resume=resume,
        strengths="Communication, Problem-solving, Teamwork",
        areas_to_improve="Technical knowledge, Leadership, Time management",
        ats_score=random.randint(60, 98)  # Random ATS score between 60 and 98
    )
    
    # Create mock skill matches
    skill_names = ["Python", "Java", "JavaScript", "HTML", "CSS", "Communication", "Teamwork"]
    for skill_name in skill_names:
        skill, _ = Skill.objects.get_or_create(name=skill_name)
        SkillMatch.objects.create(
            analysis=analysis,
            skill=skill,
            proficiency=random.randint(50, 95)
        )
    
    # Create mock recommended skills
    recommended_skill_names = ["Django", "React", "Node.js", "Data Analysis", "Machine Learning"]
    for skill_name in recommended_skill_names:
        skill, _ = Skill.objects.get_or_create(name=skill_name)
        RecommendedSkill.objects.create(
            analysis=analysis,
            skill=skill,
            importance=random.randint(70, 95)
        )
    
    # Create mock career path recommendations
    career_paths = [
        ("Software Developer", "Develop applications and systems using programming languages.", "₹6-15 LPA", "High"),
        ("Data Analyst", "Analyze and interpret complex data to inform business decisions.", "₹5-12 LPA", "High"),
        ("Web Developer", "Create and maintain websites and web applications.", "₹4-10 LPA", "Medium"),
        ("UI/UX Designer", "Design user interfaces and experiences for digital products.", "₹4-8 LPA", "Medium"),
        ("Project Manager", "Plan, execute, and close projects within constraints.", "₹8-18 LPA", "Medium")
    ]
    
    for title, desc, salary, growth in career_paths:
        career_path, _ = CareerPath.objects.get_or_create(
            title=title,
            defaults={
                'description': desc,
                'avg_salary': salary,
                'growth_potential': growth
            }
        )
        
        CareerPathRecommendation.objects.create(
            analysis=analysis,
            career_path=career_path,
            match_percentage=random.randint(60, 95)
        )
    
    # Create mock job recommendations
    # First ensure we have job portals
    job_portals = [
        ("Naukri.com", "https://www.naukri.com"),
        ("Indeed.com", "https://www.indeed.co.in"),
        ("LinkedIn", "https://www.linkedin.com")
    ]
    
    portal_objects = {}
    for name, url in job_portals:
        portal, _ = JobPortal.objects.get_or_create(
            name=name, 
            defaults={'url': url}
        )
        portal_objects[name] = portal
    
    # Now create job recommendations
    job_recommendations = [
        ("Junior Software Developer", "TCS", "Entry-level development role", portal_objects["Naukri.com"], "https://www.naukri.com/job-listings-junior-software-developer"),
        ("Graduate Trainee", "Infosys", "Training program for fresh graduates", portal_objects["Indeed.com"], "https://www.indeed.co.in/viewjob?jk=12345"),
        ("Web Developer Intern", "Wipro", "Web development internship opportunity", portal_objects["LinkedIn"], "https://www.linkedin.com/jobs/view/12345"),
        ("Associate Data Analyst", "Cognizant", "Entry-level data analysis position", portal_objects["Naukri.com"], "https://www.naukri.com/job-listings-associate-data-analyst"),
        ("UI Designer", "Tech Mahindra", "Junior UI designer role", portal_objects["Indeed.com"], "https://www.indeed.co.in/viewjob?jk=67890")
    ]
    
    for title, company, desc, portal, url in job_recommendations:
        JobRecommendation.objects.create(
            analysis=analysis,
            title=title,
            company=company,
            description=desc,
            portal=portal,
            url=url
        )
    
    # Mark the resume as processed
    resume.processed = True
    resume.save()

# Resume Generator views
def resume_templates(request):
    """Display available resume templates"""
    return render(request, 'compass_app/resume_templates.html')

def resume_generator(request):
    """Resume generator form based on selected template"""
    template_type = request.GET.get('template', 'professional')
    
    template_names = {
        'professional': 'Professional',
        'modern': 'Modern',
        'entry-level': 'Entry Level'
    }
    
    template_name = template_names.get(template_type, 'Professional')
    
    context = {
        'template_type': template_type,
        'template_name': template_name
    }
    
    return render(request, 'compass_app/resume_generator.html', context)

def generate_resume(request):
    """Generate and download the resume based on form data"""
    if request.method != 'POST':
        return redirect('resume_templates')
    
    # Get form data
    template_type = request.POST.get('template', 'professional')
    format_type = request.POST.get('format', 'docx')
    
    # Get personal information
    full_name = request.POST.get('fullName', '')
    job_title = request.POST.get('jobTitle', '')
    email = request.POST.get('email', '')
    phone = request.POST.get('phone', '')
    location = request.POST.get('location', '')
    linkedin = request.POST.get('linkedin', '')
    summary = request.POST.get('summary', '')
    
    # Get skills
    technical_skills = request.POST.get('technicalSkills', '').split(',')
    soft_skills = request.POST.get('softSkills', '').split(',')
    
    # Create a new Word document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Apply different styling based on template type
    if template_type == 'professional':
        # Header with name and title
        name_heading = doc.add_heading('', level=0)
        name_run = name_heading.add_run(full_name.upper())
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        name_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Job title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(job_title)
        title_run.font.size = Pt(14)
        
        # Contact information
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.add_run(f"{email} | {phone} | {location}")
        if linkedin:
            contact_paragraph.add_run(f" | {linkedin}")
        
        # Summary
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        doc.add_paragraph(summary)
        
        # Skills
        doc.add_heading('SKILLS', level=1)
        skills_paragraph = doc.add_paragraph()
        skills_paragraph.add_run('Technical Skills: ').bold = True
        skills_paragraph.add_run(', '.join(s.strip() for s in technical_skills if s.strip()))
        skills_paragraph.add_run('\n')
        skills_paragraph.add_run('Soft Skills: ').bold = True
        skills_paragraph.add_run(', '.join(s.strip() for s in soft_skills if s.strip()))
        
    elif template_type == 'modern':
        # Header with name and title
        name_heading = doc.add_heading('', level=0)
        name_run = name_heading.add_run(full_name)
        name_run.font.size = Pt(22)
        name_run.font.bold = True
        
        # Job title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(job_title)
        title_run.font.size = Pt(14)
        title_run.italic = True
        
        # Contact information
        contact_paragraph = doc.add_paragraph(style='List Bullet')
        contact_paragraph.add_run(f"Email: {email}")
        doc.add_paragraph(style='List Bullet').add_run(f"Phone: {phone}")
        doc.add_paragraph(style='List Bullet').add_run(f"Location: {location}")
        if linkedin:
            doc.add_paragraph(style='List Bullet').add_run(f"LinkedIn: {linkedin}")
        
        # Summary
        doc.add_heading('PROFILE', level=1)
        doc.add_paragraph(summary)
        
        # Skills
        doc.add_heading('SKILLS', level=1)
        # Technical skills
        doc.add_paragraph().add_run('Technical').bold = True
        for skill in [s.strip() for s in technical_skills if s.strip()]:
            doc.add_paragraph(skill, style='List Bullet')
        # Soft skills
        doc.add_paragraph().add_run('Soft Skills').bold = True
        for skill in [s.strip() for s in soft_skills if s.strip()]:
            doc.add_paragraph(skill, style='List Bullet')
            
    else:  # entry-level template
        # Header with name and title
        name_heading = doc.add_heading('', level=0)
        name_run = name_heading.add_run(full_name)
        name_run.font.size = Pt(20)
        name_run.font.bold = True
        
        # Contact information in a single line
        contact_paragraph = doc.add_paragraph()
        contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_paragraph.add_run(f"{email} | {phone} | {location}")
        if linkedin:
            contact_paragraph.add_run(f" | {linkedin}")
        
        # Job title
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(job_title)
        title_run.font.size = Pt(14)
        title_run.bold = True
        
        # Summary
        doc.add_heading('CAREER OBJECTIVE', level=1)
        doc.add_paragraph(summary)
        
        # Skills in two columns
        doc.add_heading('SKILLS', level=1)
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        cell1 = table.cell(0, 0)
        cell1.text = "Technical Skills"
        cell1.paragraphs[0].runs[0].bold = True
        
        cell2 = table.cell(0, 1)
        cell2.text = "Soft Skills"
        cell2.paragraphs[0].runs[0].bold = True
        
        # Add a row for the skills
        row = table.add_row()
        cell1 = row.cells[0]
        for skill in [s.strip() for s in technical_skills if s.strip()]:
            p = cell1.add_paragraph(skill, style='List Bullet')
            
        cell2 = row.cells[1]
        for skill in [s.strip() for s in soft_skills if s.strip()]:
            p = cell2.add_paragraph(skill, style='List Bullet')
    
    # Common sections for all templates
    
    # Education
    doc.add_heading('EDUCATION', level=1)
    education_count = 0
    while True:
        degree = request.POST.get(f'education[{education_count}][degree]', '')
        institution = request.POST.get(f'education[{education_count}][institution]', '')
        start_date = request.POST.get(f'education[{education_count}][startDate]', '')
        end_date = request.POST.get(f'education[{education_count}][endDate]', '')
        description = request.POST.get(f'education[{education_count}][description]', '')
        
        if not degree or not institution:
            break
            
        # Format dates
        if start_date:
            start_date = "-".join(start_date.split("-"))
        if end_date:
            end_date = "-".join(end_date.split("-"))
            
        date_range = f"{start_date} to {end_date}" if start_date and end_date else ""
        
        p = doc.add_paragraph()
        p.add_run(f"{degree}, {institution}").bold = True
        if date_range:
            p.add_run(f" ({date_range})")
        if description:
            doc.add_paragraph(description)
            
        education_count += 1
    
    # Experience
    doc.add_heading('WORK EXPERIENCE', level=1)
    experience_count = 0
    while True:
        title = request.POST.get(f'experience[{experience_count}][title]', '')
        company = request.POST.get(f'experience[{experience_count}][company]', '')
        start_date = request.POST.get(f'experience[{experience_count}][startDate]', '')
        end_date = request.POST.get(f'experience[{experience_count}][endDate]', '')
        is_current = request.POST.get(f'experience[{experience_count}][current]', '') == 'on'
        description = request.POST.get(f'experience[{experience_count}][description]', '')
        
        if not title or not company:
            break
            
        # Format dates
        if start_date:
            start_date = "-".join(start_date.split("-"))
        if end_date and not is_current:
            end_date = "-".join(end_date.split("-"))
        else:
            end_date = "Present"
            
        date_range = f"{start_date} to {end_date}" if start_date else ""
        
        p = doc.add_paragraph()
        p.add_run(f"{title} - {company}").bold = True
        if date_range:
            p.add_run(f" ({date_range})")
        
        if description:
            # Split by newlines and add as bullet points
            for line in description.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip(), style='List Bullet')
            
        experience_count += 1
    
    # Projects (optional)
    project_count = 0
    has_projects = False
    while True:
        name = request.POST.get(f'projects[{project_count}][name]', '')
        url = request.POST.get(f'projects[{project_count}][url]', '')
        description = request.POST.get(f'projects[{project_count}][description]', '')
        
        if not name:
            break
            
        if not has_projects:
            doc.add_heading('PROJECTS', level=1)
            has_projects = True
            
        p = doc.add_paragraph()
        p.add_run(name).bold = True
        if url:
            p.add_run(f" - {url}")
        
        if description:
            doc.add_paragraph(description)
            
        project_count += 1
    
    # Certifications (optional)
    cert_count = 0
    has_certs = False
    while True:
        name = request.POST.get(f'certifications[{cert_count}][name]', '')
        issuer = request.POST.get(f'certifications[{cert_count}][issuer]', '')
        date = request.POST.get(f'certifications[{cert_count}][date]', '')
        
        if not name:
            break
            
        if not has_certs:
            doc.add_heading('CERTIFICATIONS', level=1)
            has_certs = True
            
        p = doc.add_paragraph()
        p.add_run(name).bold = True
        if issuer:
            p.add_run(f" - {issuer}")
        if date:
            p.add_run(f" ({'-'.join(date.split('-'))})")
            
        cert_count += 1
    
    # Save to a temporary file
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    temp_file.close()
    
    # Convert to PDF if needed
    if format_type == 'pdf':
        # First check if Word is installed
        if not is_word_installed():
            messages.warning(request, "Microsoft Word is not detected on the server. Providing DOCX file instead.")
            with open(temp_file.name, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{full_name.replace(" ", "_")}_Resume.docx"'
            os.unlink(temp_file.name)
            return response
            
        pdf_temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        pdf_temp_file.close()
        try:
            # Log conversion attempt
            print(f"Attempting to convert {temp_file.name} to {pdf_temp_file.name}")
            convert(temp_file.name, pdf_temp_file.name)
            
            # Create the HTTP response with the appropriate headers for PDF
            with open(pdf_temp_file.name, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/pdf')
                response['Content-Disposition'] = f'attachment; filename="{full_name.replace(" ", "_")}_Resume.pdf"'
            
            # Delete the temporary files
            os.unlink(temp_file.name)
            os.unlink(pdf_temp_file.name)
            
        except FileNotFoundError as e:
            # Most likely Microsoft Word is not installed or accessible
            messages.warning(request, f"PDF conversion failed: Microsoft Word not found. Providing DOCX file instead. Error: {str(e)}")
            with open(temp_file.name, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{full_name.replace(" ", "_")}_Resume.docx"'
            os.unlink(temp_file.name)
        except PermissionError as e:
            # File access issues
            messages.warning(request, f"PDF conversion failed: Permission error. Providing DOCX file instead. Error: {str(e)}")
            with open(temp_file.name, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{full_name.replace(" ", "_")}_Resume.docx"'
            os.unlink(temp_file.name)
        except Exception as e:
            # If conversion fails for any other reason, fall back to DOCX
            error_message = str(e)
            messages.warning(request, f"PDF conversion failed: {error_message}. Providing DOCX file instead.")
            with open(temp_file.name, 'rb') as file:
                response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                response['Content-Disposition'] = f'attachment; filename="{full_name.replace(" ", "_")}_Resume.docx"'
            os.unlink(temp_file.name)
    else:
        # Create the HTTP response with the appropriate headers for DOCX
        with open(temp_file.name, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{full_name.replace(" ", "_")}_Resume.docx"'
        
        # Delete the temporary file
        os.unlink(temp_file.name)
    
    # Add a success message
    messages.success(request, f"Resume successfully generated and downloaded as {format_type.upper()}!")
    
    return response

def feedback(request):
    context = {}
    
    if request.method == 'POST':
        # Get form data
        name = request.POST.get('name', '')
        email = request.POST.get('email', '')
        rating = request.POST.get('rating', '')
        feedback_text = request.POST.get('feedback', '')
        contact_permission = request.POST.get('contact_permission') == 'on'
        
        # Create email subject and message
        subject = f"Career Compass Feedback - {rating} Stars from {name}"
        
        # Format the message body
        message = f"""
Feedback received from Career Compass user:

Name: {name}
Email: {email}
Rating: {rating} out of 5 stars
Allow Contact: {'Yes' if contact_permission else 'No'}

Feedback:
{feedback_text}
        """
        
        # Send email to administrator
        try:
            send_mail(
                subject,
                message,
                settings.DEFAULT_FROM_EMAIL,  # Use site's default email as sender
                [settings.ADMIN_EMAIL],  # Send to admin email specified in settings
                fail_silently=False,
            )
            context['success_message'] = "Thank you for your feedback! We appreciate your time."
        except Exception as e:
            messages.error(request, f"There was an error sending your feedback: {str(e)}")
    
    return render(request, 'compass_app/feedback.html', context)

def remove_resume(request):
    """Remove the user's resume and clear session data."""
    if request.method == 'POST':
        # Remove resume from DB if it exists (for demo, just clear session)
        request.session.pop('resume_analysis', None)
        request.session.pop('resume_processed', None)
        request.session.pop('resume_filename', None)
        request.session.pop('resume_uploaded_at', None)
        messages.success(request, 'Your resume has been removed from the system.')
    return redirect('dashboard')

def mock_interview(request):
    """Provide a mock interview experience powered by AI."""
    # Check if resume has been processed from session
    resume_processed = request.session.get('resume_processed', False)
    resume_analysis = request.session.get('resume_analysis', {})
    
    if not resume_processed:
        messages.warning(request, 'Please upload your resume first to get a personalized mock interview experience.')
        return redirect('resume_upload')
    
    # Extract relevant information from resume analysis
    skills = resume_analysis.get('skills', [])
    experience = resume_analysis.get('experience', [])
    education = resume_analysis.get('education', [])
    job_title = resume_analysis.get('job_title', 'Job Applicant')
    
    # Context to store all interview-related data
    context = {
        'resume_processed': resume_processed,
        'job_title': job_title,
        'interview_started': False,
        'interview_completed': False,
        'current_question': 0,
        'evaluation': None,
        'final_score': None
    }
    
    # Handle interview interactions
    if request.method == 'POST':
        action = request.POST.get('action', '')
        
        if action == 'start_interview':
            # Generate interview questions based on resume data
            interview_questions = generate_interview_questions(skills, experience, education, job_title)
            request.session['interview_questions'] = interview_questions
            request.session['current_question'] = 0
            request.session['answers'] = []
            request.session['evaluations'] = []
            request.session['interview_started'] = True
            request.session['interview_completed'] = False
            
            context['interview_started'] = True
            context['interview_questions'] = interview_questions
            context['current_question'] = 0
            
        elif action == 'submit_answer':
            answer = request.POST.get('answer', '')
            current_question = int(request.POST.get('current_question', 0))
            interview_questions = request.session.get('interview_questions', [])
            
            if current_question < len(interview_questions):
                # Get existing answers and evaluations or initialize empty lists
                answers = request.session.get('answers', [])
                evaluations = request.session.get('evaluations', [])
                
                # Evaluate answer
                evaluation = evaluate_answer(answer, interview_questions[current_question], skills)
                
                # Store answer and evaluation
                answers.append(answer)
                evaluations.append(evaluation)
                request.session['answers'] = answers
                request.session['evaluations'] = evaluations
                
                # Move to next question or complete interview
                next_question = current_question + 1
                request.session['current_question'] = next_question
                
                if next_question >= len(interview_questions):
                    # Interview completed
                    final_score = calculate_final_score(evaluations)
                    request.session['final_score'] = final_score
                    request.session['interview_completed'] = True
                    
                    context['interview_completed'] = True
                    context['final_score'] = final_score
                    context['evaluations'] = evaluations
                else:
                    context['current_question'] = next_question
            
            context['interview_started'] = True
            context['interview_questions'] = interview_questions
            context['answers'] = request.session.get('answers', [])
            context['evaluations'] = request.session.get('evaluations', [])
            
        elif action == 'restart_interview':
            # Clear interview session data
            request.session.pop('interview_questions', None)
            request.session.pop('current_question', None)
            request.session.pop('answers', None)
            request.session.pop('evaluations', None)
            request.session.pop('interview_started', None)
            request.session.pop('interview_completed', None)
            request.session.pop('final_score', None)
            
            # Redirect to the interview page to start fresh
            return redirect('mock_interview')
    else:
        # For GET requests, check if an interview is already in progress
        if request.session.get('interview_started', False):
            context['interview_started'] = True
            context['interview_questions'] = request.session.get('interview_questions', [])
            context['current_question'] = request.session.get('current_question', 0)
            context['answers'] = request.session.get('answers', [])
            context['evaluations'] = request.session.get('evaluations', [])
            context['interview_completed'] = request.session.get('interview_completed', False)
            context['final_score'] = request.session.get('final_score', None)
    
    return render(request, 'compass_app/mock_interview.html', context)

def generate_interview_questions(skills, experience, education, job_title):
    """Generate relevant interview questions based on resume data."""
    # Define standard questions that work for all candidates
    standard_questions = [
        "Tell me about yourself and your background.",
        "What are your greatest professional strengths?",
        "What do you consider to be your weaknesses?",
        "Why are you interested in this position?",
        "Where do you see yourself in 5 years?"
    ]
    
    # Skills-based questions
    skill_questions = []
    if skills:
        # Technical skills questions
        technical_skills = [skill for skill in skills if skill.lower() in [
            'python', 'java', 'javascript', 'html', 'css', 'react', 'angular', 'node.js',
            'django', 'flask', 'express', 'sql', 'mysql', 'postgresql', 'mongodb',
            'git', 'docker', 'aws', 'azure', 'gcp', 'linux', 'windows', 'macos',
            'c++', 'c#', 'ruby', 'php', 'go', 'swift', 'kotlin', 'typescript',
            '.net', 'spring', 'hibernate', 'rest api', 'graphql', 'json', 'xml', 'yaml',
            'nosql', 'jenkins', 'kubernetes', 'ci/cd', 'machine learning', 'data science',
            'tensorflow', 'pytorch', 'nlp', 'computer vision', 'artificial intelligence'
        ]]
        
        # Soft skills questions
        soft_skills = [skill for skill in skills if skill.lower() in [
            'communication', 'teamwork', 'leadership', 'problem solving', 'critical thinking',
            'time management', 'project management', 'agile', 'scrum', 'devops'
        ]]
        
        # Generate questions for technical skills
        for skill in technical_skills[:3]:  # Limit to top 3 technical skills
            skill_questions.extend([
                f"Can you describe a project where you utilized your {skill} skills?",
                f"How do you stay updated with the latest developments in {skill}?",
                f"What challenges have you faced while working with {skill} and how did you overcome them?",
                f"Can you explain a complex {skill} concept to someone with limited technical knowledge?",
                f"How do you ensure the quality of your {skill} code?"
            ])
        
        # Generate questions for soft skills
        for skill in soft_skills[:2]:  # Limit to top 2 soft skills
            skill_questions.extend([
                f"Can you give an example of how you've demonstrated {skill} in a professional setting?",
                f"How do you maintain and improve your {skill}?",
                f"Describe a situation where your {skill} helped resolve a conflict or problem.",
                f"How do you measure the effectiveness of your {skill}?"
            ])
    
    # Experience-based questions
    experience_questions = []
    if experience:
        experience_questions.extend([
            "Describe a challenging situation you faced at work and how you handled it.",
            "Tell me about a time when you had to work under pressure to meet a deadline.",
            "Give an example of a goal you reached and how you achieved it.",
            "Describe a time when you had to learn a new technology quickly.",
            "Tell me about a project where you had to collaborate with a difficult team member."
        ])
    
    # Education-based questions
    education_questions = []
    if education:
        education_questions.extend([
            "How has your education prepared you for this career?",
            "What extracurricular activities were you involved in during your studies?",
            "How do you apply what you learned in your education to your current work?",
            "What was your favorite subject and why?",
            "How do you continue to learn and grow professionally?"
        ])
    
    # Job-specific questions
    job_specific_questions = [
        f"What attracted you to apply for the {job_title} position?",
        f"What unique skills would you bring to the role of {job_title}?",
        f"What do you know about the day-to-day responsibilities of a {job_title}?",
        f"How do your skills align with the requirements of a {job_title}?",
        f"What challenges do you anticipate in the role of {job_title} and how would you handle them?"
    ]
    
    # Combine all questions and select 10 (or less if not enough)
    all_questions = standard_questions + skill_questions + experience_questions + education_questions + job_specific_questions
    
    # Ensure we have at least 10 questions by adding generic ones if needed
    additional_questions = [
        "How do you handle stress and pressure?",
        "How would your colleagues describe you?",
        "What is your preferred work environment?",
        "How do you prioritize your work?",
        "What are you passionate about?",
        "Describe your ideal manager.",
        "What motivates you to perform well?",
        "How do you handle criticism?",
        "What are your salary expectations?",
        "Do you have any questions for me?"
    ]
    
    while len(all_questions) < 10:
        all_questions.append(additional_questions.pop(0))
    
    # Shuffle and select 10 questions
    import random
    random.shuffle(all_questions)
    return all_questions[:10]

def evaluate_answer(answer, question, skills):
    """Evaluate the interview answer and provide feedback."""
    # Check if answer is substantial (arbitrary minimum length)
    if len(answer.strip()) < 30:
        score = 2
        feedback = "Your answer is too brief. Try to provide more details and examples to demonstrate your experience and skills."
        improvement = "Elaborate on your points and provide specific examples from your past experience."
    elif len(answer.strip()) < 100:
        score = 6
        feedback = "Your answer is good but could benefit from more depth and specific examples."
        improvement = "Include more concrete examples and quantify achievements when possible."
    else:
        # Check for relevant keywords related to the question
        relevance_score = check_answer_relevance(answer, question, skills)
        
        if relevance_score < 0.3:
            score = 4
            feedback = "Your answer seems somewhat off-topic. Try to focus more directly on addressing the question."
            improvement = "Make sure you understand the question fully before answering and stay focused on the key points."
        elif relevance_score < 0.6:
            score = 7
            feedback = "Good answer that addresses the question, but could be more focused on key points."
            improvement = "Structure your answer with a clear introduction, key points, and conclusion."
        else:
            score = 9
            feedback = "Excellent answer that thoroughly addresses the question with good examples."
            improvement = "Continue to refine your delivery and consider preparing even more diverse examples."
    
    # Check for confidence markers in language
    confidence_score = check_confidence_markers(answer)
    
    # Generate body language tip
    body_language_tip = generate_body_language_tip()
    
    # Generate suggested answer based on question type
    suggested_answer = generate_suggested_answer(question, skills)
    
    # Final evaluation object
    evaluation = {
        'score': score,
        'feedback': feedback,
        'improvement': improvement,
        'confidence_score': confidence_score,
        'body_language_tip': body_language_tip,
        'suggested_answer': suggested_answer
    }
    
    return evaluation

def check_answer_relevance(answer, question, skills):
    """Check how relevant the answer is to the question asked."""
    answer = answer.lower()
    question = question.lower()
    
    # Extract key terms from the question
    import re
    key_terms = re.findall(r'\b\w{3,}\b', question)
    key_terms = [term for term in key_terms if term not in ('the', 'and', 'for', 'you', 'your', 'about', 'what', 'how', 'why', 'when', 'where', 'who')]
    
    # Add skills to key terms if the question mentions skills
    if 'skill' in question or 'ability' in question or 'expertise' in question:
        key_terms.extend([s.lower() for s in skills])
    
    # Count how many key terms are in the answer
    term_count = sum(1 for term in key_terms if term in answer)
    
    # Calculate relevance score (0 to 1)
    if not key_terms:
        return 0.5  # Default middle value if no key terms
    
    return min(1.0, term_count / len(key_terms))

def check_confidence_markers(answer):
    """Check for confidence or uncertainty in language."""
    confidence_markers = ['definitely', 'certainly', 'absolutely', 'confident', 'sure', 'expert', 'specialized', 'experienced']
    uncertainty_markers = ['maybe', 'perhaps', 'i think', 'possibly', 'might', 'not sure', 'i guess', 'kind of', 'sort of']
    
    answer = answer.lower()
    
    confidence_count = sum(1 for marker in confidence_markers if marker in answer)
    uncertainty_count = sum(1 for marker in uncertainty_markers if marker in answer)
    
    # Calculate confidence score (0 to 1)
    total_markers = confidence_count + uncertainty_count
    if total_markers == 0:
        return 0.5  # Default middle value
    
    return confidence_count / total_markers

def generate_body_language_tip():
    """Generate a random body language tip for interviews."""
    tips = [
        "Maintain good eye contact to convey confidence and engagement.",
        "Sit with a straight posture to project confidence and attentiveness.",
        "Use hand gestures moderately to emphasize points, but avoid excessive movements.",
        "Smile naturally at appropriate moments to establish rapport.",
        "Mirror the interviewer's body language subtly to build connection.",
        "Avoid fidgeting or playing with objects as it signals nervousness.",
        "Lean slightly forward to show interest and engagement in the conversation.",
        "Nod occasionally to show active listening and agreement.",
        "Keep your hands visible, not hidden under the table.",
        "Take a moment to breathe and collect your thoughts before answering difficult questions."
    ]
    
    import random
    return random.choice(tips)

def calculate_final_score(evaluations):
    """Calculate the final interview score based on all question evaluations."""
    if not evaluations:
        return 0
    
    # Extract scores from each evaluation
    scores = [eval.get('score', 0) for eval in evaluations]
    
    # Calculate average score (0-10 scale)
    avg_score = sum(scores) / len(scores)
    
    return round(avg_score, 1)

def generate_suggested_answer(question, skills):
    """Generate a suggested best answer for the given question."""
    question = question.lower()
    
    # Standard questions
    if "tell me about yourself" in question:
        return """I would suggest structuring your answer in this way:
1. Start with your current role and key responsibilities
2. Highlight your relevant experience and achievements
3. Mention your education and key skills
4. Conclude with your career goals and interest in the position

Example: 'I am currently working as a [Your Role] with [X] years of experience in [Industry]. I specialize in [Key Skills] and have successfully [Key Achievement]. I hold a [Degree] from [University] and am particularly interested in this role because [Specific Reason].'"""
    
    elif "greatest professional strengths" in question:
        return """Focus on 2-3 key strengths that are relevant to the position:
1. Choose strengths that align with the job requirements
2. Provide specific examples for each strength
3. Quantify your achievements if possible

Example: 'My greatest strengths are [Strength 1] and [Strength 2]. For example, in my previous role, I [Specific Example]. This resulted in [Quantifiable Result].'"""
    
    elif "weaknesses" in question:
        return """Structure your answer using this approach:
1. Choose a real but not critical weakness
2. Show how you're working to improve it
3. Demonstrate self-awareness and growth mindset

Example: 'One area I'm working to improve is [Weakness]. I've been addressing this by [Action Taken], and I've seen improvement in [Specific Result].'"""
    
    elif "why are you interested" in question:
        return """Include these key elements:
1. Your interest in the company/industry
2. How your skills match the role
3. Your career goals and how they align
4. Specific aspects of the position that excite you

Example: 'I'm particularly interested in this position because [Company/Industry Reason]. My experience in [Relevant Skill] and [Another Skill] aligns well with the requirements. I'm excited about the opportunity to [Specific Aspect of Role].'"""
    
    # Technical skills questions
    elif any(skill.lower() in question for skill in skills if skill.lower() in [
        'python', 'java', 'javascript', 'html', 'css', 'react', 'angular', 'node.js',
        'django', 'flask', 'express', 'sql', 'mysql', 'postgresql', 'mongodb',
        'git', 'docker', 'aws', 'azure', 'gcp', 'linux', 'windows', 'macos',
        'c++', 'c#', 'ruby', 'php', 'go', 'swift', 'kotlin', 'typescript',
        '.net', 'spring', 'hibernate', 'rest api', 'graphql', 'json', 'xml', 'yaml',
        'nosql', 'jenkins', 'kubernetes', 'ci/cd', 'machine learning', 'data science',
        'tensorflow', 'pytorch', 'nlp', 'computer vision', 'artificial intelligence'
    ]):
        return """For technical skill questions, follow this structure:
1. Describe a specific project where you used the skill
2. Explain the challenges you faced
3. Detail your approach and solution
4. Share the results and what you learned

Example: 'In my previous project, I used [Skill] to [Project Goal]. The main challenge was [Challenge]. I addressed this by [Solution]. The result was [Outcome], and I learned [Key Learning].'"""
    
    # Soft skills questions
    elif any(skill.lower() in question for skill in skills if skill.lower() in [
        'communication', 'teamwork', 'leadership', 'problem solving', 'critical thinking',
        'time management', 'project management', 'agile', 'scrum', 'devops'
    ]):
        return """For soft skill questions, use the STAR method:
1. Situation: Describe the context
2. Task: Explain what needed to be done
3. Action: Detail what you did
4. Result: Share the outcome

Example: 'In a situation where [Context], I needed to [Task]. I demonstrated [Skill] by [Action]. This resulted in [Positive Outcome].'"""
    
    # Experience-based questions
    elif any(term in question for term in ['challenging situation', 'pressure', 'deadline', 'goal', 'learn']):
        return """For experience-based questions, use this framework:
1. Set the context and challenge
2. Explain your thought process
3. Describe your actions
4. Share the results and lessons learned

Example: 'When faced with [Challenge], I first [Thought Process]. I then [Actions Taken]. The outcome was [Result], and I learned [Key Lessons].'"""
    
    # Education-based questions
    elif any(term in question for term in ['education', 'studies', 'learn', 'subject']):
        return """For education questions, focus on:
1. Relevant coursework and projects
2. Practical applications of your learning
3. How your education prepared you for this role
4. Your approach to continuous learning

Example: 'My education in [Field] provided me with [Key Knowledge]. I applied this in [Project/Experience]. This has prepared me for [Aspect of Role] and I continue to learn through [Current Learning Activities].'"""
    
    # Default suggested answer
    return """For this question, consider:
1. Understanding what the interviewer is really asking
2. Providing specific examples from your experience
3. Showing how your skills and experience match the role
4. Being concise but thorough in your response

Remember to:
- Stay focused on the question
- Use specific examples
- Show enthusiasm and confidence
- Be honest and authentic"""

def interview_practice(request):
    """View for interview practice with indiabix.com integration"""
    
    # Get the topic categories from indiabix
    categories = [
        {'name': 'Aptitude', 'url': 'https://www.indiabix.com/aptitude/questions-and-answers/'},
        {'name': 'Technical Interview', 'url': 'https://www.indiabix.com/technical/interview-questions-and-answers/'},
        {'name': 'Verbal Ability', 'url': 'https://www.indiabix.com/verbal-ability/questions-and-answers/'},
        {'name': 'Logical Reasoning', 'url': 'https://www.indiabix.com/logical-reasoning/questions-and-answers/'},
        {'name': 'Data Interpretation', 'url': 'https://www.indiabix.com/data-interpretation/questions-and-answers/'},
        {'name': 'C Programming', 'url': 'https://www.indiabix.com/c-programming/questions-and-answers/'},
        {'name': 'Java Programming', 'url': 'https://www.indiabix.com/java-programming/questions-and-answers/'},
        {'name': 'Database', 'url': 'https://www.indiabix.com/database/questions-and-answers/'},
        {'name': 'HR Interview', 'url': 'https://www.indiabix.com/hr-interview/questions-and-answers/'}
    ]
    
    # Calculate suggested categories based on user's profile/skills if available
    suggested_categories = []
    
    # If resume has been processed, suggest relevant categories
    resume_processed = request.session.get('resume_processed', False)
    if resume_processed:
        analysis_data = request.session.get('resume_analysis', {})
        skills = analysis_data.get('skill_matches', [])
        career_paths = analysis_data.get('career_recommendations', [])
        
        # Map skills to relevant categories
        skill_to_category = {
            'programming': ['C Programming', 'Java Programming'],
            'database': ['Database'],
            'communication': ['Verbal Ability', 'HR Interview'],
            'analytical': ['Data Interpretation', 'Logical Reasoning', 'Aptitude'],
            'problem solving': ['Logical Reasoning', 'Aptitude'],
        }
        
        # Add suggested categories based on user's skills
        for skill in skills:
            skill_name = skill.get('name', '').lower()
            for key, categories_list in skill_to_category.items():
                if key in skill_name:
                    for category_name in categories_list:
                        for category in categories:
                            if category['name'] == category_name and category not in suggested_categories:
                                suggested_categories.append(category)
    
    context = {
        'categories': categories,
        'suggested_categories': suggested_categories
    }
    
    return render(request, 'compass_app/interview_practice.html', context)
